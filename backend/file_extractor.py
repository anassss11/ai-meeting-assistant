"""
File extraction utilities for transcript uploads.
Supports: PDF, TXT, CSV, DOCX
"""

import logging
from pathlib import Path
from typing import Optional
import PyPDF2
import pandas as pd
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        Extracted text
    """
    try:
        text = []
        with open(file_path, 'rb') as pdf_file:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        
        result = '\n'.join(text)
        logger.info(f"Extracted {len(result)} characters from PDF")
        return result
    
    except Exception as e:
        logger.error(f"Error extracting PDF: {e}")
        raise


def extract_text_from_txt(file_path: Path) -> str:
    """
    Extract text from TXT file.
    
    Args:
        file_path: Path to TXT file
    
    Returns:
        Extracted text
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as txt_file:
            text = txt_file.read()
        
        logger.info(f"Extracted {len(text)} characters from TXT")
        return text
    
    except Exception as e:
        logger.error(f"Error extracting TXT: {e}")
        raise


def extract_text_from_csv(file_path: Path) -> str:
    """
    Extract text from CSV file.
    
    Args:
        file_path: Path to CSV file
    
    Returns:
        Extracted text (CSV content as formatted string)
    """
    try:
        df = pd.read_csv(file_path)
        
        # Convert DataFrame to readable text format
        text = df.to_string()
        
        logger.info(f"Extracted {len(text)} characters from CSV")
        return text
    
    except Exception as e:
        logger.error(f"Error extracting CSV: {e}")
        raise


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from DOCX (Word) file.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        Extracted text
    """
    try:
        doc = Document(file_path)
        
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        text = "\n".join(text_parts)
        
        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text
    
    except Exception as e:
        logger.error(f"Error extracting DOCX: {e}")
        raise


def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text from file based on extension.
    
    Args:
        file_path: Path to file
    
    Returns:
        Extracted text
    
    Raises:
        ValueError: If file type is not supported
    """
    suffix = file_path.suffix.lower()
    
    if suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix == '.txt':
        return extract_text_from_txt(file_path)
    elif suffix == '.csv':
        return extract_text_from_csv(file_path)
    elif suffix == '.docx':
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")
